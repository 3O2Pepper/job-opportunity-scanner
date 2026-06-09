"""Parse resumes and profile text; maintain merged profile snapshot."""

from __future__ import annotations

import hashlib
import json
import re
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy import select

from app.db.models import ProfileDocument, ProfileSnapshot
from app.db.session import session_scope


def _sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()


def extract_text_from_docx(data: bytes) -> str:
    doc = DocxDocument(BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs).strip()


def extract_text_from_upload(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(data)
    if name.endswith(".docx"):
        return extract_text_from_docx(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="replace").strip()
    raise ValueError(f"Unsupported file type: {filename}")


_SIMPLE_SKILL_SPLIT = re.compile(r"[,;/•\n]")


def infer_skills_from_text(text: str, max_skills: int = 80) -> list[str]:
    """Very simple skill tokenization for deterministic scoring."""
    lower = text.lower()
    found: set[str] = set()
    # Common resume section heuristic
    for line in lower.splitlines():
        line = line.strip()
        if not line:
            continue
        if any(h in line for h in ("skills", "tools", "software", "technologies")):
            for chunk in _SIMPLE_SKILL_SPLIT.split(line):
                chunk = chunk.strip()
                if 2 <= len(chunk) <= 40:
                    found.add(chunk)
    # Fallback: keywords that look like tools (caps / mixed)
    for word in re.findall(r"\b[A-Z]{2,10}\b", text):
        found.add(word.lower())
    return sorted(found)[:max_skills]


def infer_years_experience_band(text: str) -> str:
    t = text.lower()
    if any(x in t for x in ("intern", "co-op", "coop", "student")):
        return "intern"
    if any(x in t for x in ("new grad", "recent graduate", "entry level")):
        return "entry"
    m = re.search(r"(\d+)\s*\+\s*years", t)
    if m and int(m.group(1)) <= 2:
        return "early"
    return "general"


def extract_career_goals_text(docs: list[ProfileDocument]) -> str:
    """Return the concatenated text of all ``preferences`` documents."""
    parts = [d.content_text for d in docs if d.kind == "preferences"]
    return "\n".join(parts).strip()


def build_structured_profile(merged_text: str, career_goals: str = "") -> dict:
    return {
        "skills": infer_skills_from_text(merged_text),
        "experience_band": infer_years_experience_band(merged_text),
        "raw_length": len(merged_text),
        "career_goals": career_goals,
    }


def merge_profile_documents(docs: list[ProfileDocument]) -> str:
    blocks: list[str] = []
    for d in sorted(docs, key=lambda x: (x.kind, x.id)):
        label = d.kind
        if d.title:
            label = f"{d.kind}: {d.title}"
        blocks.append(f"### {label}\n{d.content_text.strip()}")
    return "\n\n".join(blocks).strip()


def refresh_profile_snapshot(session) -> ProfileSnapshot:
    docs = session.scalars(select(ProfileDocument).order_by(ProfileDocument.id.asc())).all()
    merged = merge_profile_documents(docs) if docs else ""
    career_goals = extract_career_goals_text(docs) if docs else ""
    structured = build_structured_profile(merged, career_goals=career_goals) if merged else {}
    snap = session.scalar(
        select(ProfileSnapshot).order_by(ProfileSnapshot.id.desc()).limit(1)
    )
    if snap is None:
        snap = ProfileSnapshot(merged_text=merged, structured_json=json.dumps(structured))
        session.add(snap)
    else:
        snap.merged_text = merged
        snap.structured_json = json.dumps(structured)
    session.commit()
    session.refresh(snap)
    return snap


def save_profile_text(kind: str, content: str, title: str | None = None) -> ProfileDocument:
    content = content.strip()
    if not content:
        raise ValueError("Empty content")
    h = _sha256(content)
    with session_scope() as session:
        existing = session.scalar(
            select(ProfileDocument)
            .where(
                ProfileDocument.content_hash == h,
                ProfileDocument.kind == kind,
            )
            .limit(1)
        )
        if existing:
            refresh_profile_snapshot(session)
            return existing
        doc = ProfileDocument(kind=kind, title=title, content_text=content, content_hash=h)
        session.add(doc)
        session.commit()
        session.refresh(doc)
        refresh_profile_snapshot(session)
        return doc


def save_profile_file(filename: str, data: bytes) -> ProfileDocument:
    """Save an uploaded resume (PDF/DOCX/TXT) as ``resume_file`` kind."""
    text = extract_text_from_upload(filename, data)
    title = Path(filename).name
    return save_profile_text("resume_file", text, title=title)


def save_linkedin_file(filename: str, data: bytes) -> ProfileDocument:
    """Save an uploaded LinkedIn PDF/export as ``linkedin_export`` kind."""
    text = extract_text_from_upload(filename, data)
    title = Path(filename).name
    return save_profile_text("linkedin_export", text, title=title)


def save_career_goals(goals_text: str) -> ProfileDocument:
    """Save career goals and preferences as ``preferences`` kind."""
    return save_profile_text("preferences", goals_text.strip())


def list_profile_documents() -> list[dict]:
    """Return a UI-friendly summary of every stored profile document."""
    with session_scope() as session:
        docs = session.scalars(
            select(ProfileDocument).order_by(ProfileDocument.created_at.desc())
        ).all()
        return [
            {
                "id": d.id,
                "kind": d.kind,
                "title": d.title or "",
                "length": len(d.content_text or ""),
                "created_at": d.created_at,
                "hash_short": (d.content_hash or "")[:10],
                "preview": (d.content_text or "")[:240],
            }
            for d in docs
        ]


def delete_profile_document(doc_id: int) -> bool:
    """Delete a single ProfileDocument and refresh the snapshot. Returns True if deleted."""
    with session_scope() as session:
        doc = session.get(ProfileDocument, doc_id)
        if doc is None:
            return False
        session.delete(doc)
        session.commit()
        refresh_profile_snapshot(session)
        return True


def current_snapshot_summary() -> dict | None:
    """Return a small dict describing the latest ProfileSnapshot, or None."""
    with session_scope() as session:
        snap = session.scalar(
            select(ProfileSnapshot).order_by(ProfileSnapshot.id.desc()).limit(1)
        )
        if snap is None:
            return None
        structured: dict = {}
        if snap.structured_json:
            try:
                structured = json.loads(snap.structured_json)
            except json.JSONDecodeError:
                structured = {}
        return {
            "id": snap.id,
            "updated_at": snap.updated_at,
            "merged_length": len(snap.merged_text or ""),
            "raw_length": int(structured.get("raw_length") or 0),
            "experience_band": str(structured.get("experience_band") or "general"),
            "skills": list(structured.get("skills") or []),
            "career_goals": str(structured.get("career_goals") or ""),
            "merged_preview": (snap.merged_text or "")[:600],
        }
